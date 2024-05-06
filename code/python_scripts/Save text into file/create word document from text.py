from docx import Document

# Create a new Document
doc = Document()
doc.add_heading("Interview Questions and Answers", 0)

# Add sections and content
sections = [
    ("Role and Responsibilities", [
        ("How would you describe your role in the most recent job?",
         "In my most recent role as a Software Architect Developer Tech Lead at Bank of America, I was responsible for designing and implementing full-stack solutions. "
         "I developed an Angular 15-based web interface for real-time server component status management and created scalable Spring Boot back-end services. "
         "I also worked on log analysis and data consolidation, Ansible automation, and led a team for deployment, configuration, and modernization of applications across multiple servers.")
    ]),
    
    ("Java Topics", [
        
        ("What data structures and collections are you familiar with?", 
         "I am familiar with the following data structures and collections in Java:\n"
         "1. List:\n"
         "   - ArrayList - dynamically sized array.\n"
         "   - LinkedList - doubly linked list.\n\n"
         "2. Set:\n"
         "   - HashSet - backed by a hash table.\n"
         "   - LinkedHashSet - maintains insertion order.\n"
         "   - TreeSet - sorted set.\n\n"
         "3. Map:\n"
         "   - HashMap - hash table-based implementation.\n"
         "   - LinkedHashMap - maintains insertion order.\n"
         "   - TreeMap - sorted map.\n\n"
         "4. Queue:\n"
         "   - PriorityQueue - priority heap.\n"
         "   - ArrayDeque - double-ended queue.\n\n"
         "5. Others:\n"
         "   - Stack - LIFO stack.\n"
         "   - Array - fixed-size data structure."),

        ("Coding challenge - Convert Roman numerals to integers.", 
         "Here’s a Python implementation for converting Roman numerals to integers:\n\n"
         "```python\n"
         "class Solution:\n"
         "    def romanToInt(self, s: str) -> int:\n"
         "        roman_to_int = {\n"
         "            'I': 1,\n"
         "            'V': 5,\n"
         "            'X': 10,\n"
         "            'L': 50,\n"
         "            'C': 100,\n"
         "            'D': 500,\n"
         "            'M': 1000\n"
         "        }\n"
         "        \n"
         "        result = 0\n"
         "        prev_value = 0\n"
         "        \n"
         "        for char in s[::-1]:\n"
         "            current_value = roman_to_int[char]\n"
         "            if current_value >= prev_value:\n"
         "                result += current_value\n"
         "            else:\n"
         "                result -= current_value\n"
         "            prev_value = current_value\n"
         "        \n"
         "        return result\n"
         "```")
    ]),
    

    ("Angular Topics", [
        ("How familiar are you with Angular?",
         "I am highly familiar with Angular, having developed multiple applications using Angular 10-15. I have experience with components, services, and routing. "
         "I’ve also worked with Angular Material for UI components and used TypeScript extensively for development."),

        ("What is an Angular component, and what is its purpose?",
         "An Angular component is a building block of Angular applications. It consists of:\n"
         "1. Template - defines the view.\n"
         "2. Class - defines the logic.\n"
         "3. Styles - defines the component-specific styles.\n\n"
         "A component's purpose is to encapsulate the UI and logic for a specific part of the application, providing reusability and maintainability.")
    ])
]

for section_title, section_content in sections:
    doc.add_heading(section_title, level=1)
    for question, answer in section_content:
        doc.add_heading(question, level=2)
        doc.add_paragraph(answer)

# Save the document
file_path = "/mnt/data/Interview_Questions_Answers.docx"
doc.save(file_path)
file_path
